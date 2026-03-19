from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import Pt, Inches
import shutil
import copy

class DocxGenerator:
    """
    Fills DOCX templates by replacing [[tags]] with generated content.
    Handles bullets, descriptions, environment, and skills sections.
    """
    
    def _create_bullet_paragraph(self, doc, text, font_name="Times New Roman", font_size=10):
        """
        Create a paragraph with bullet formatting.
        Uses Word's built-in bullet list by setting numPr properties.
        """
        para = doc.add_paragraph()
        
        # --- FORCED ALIGNMENT FIX ---
        # Forces the bullet to align left and creates a hanging indent for the text
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.first_line_indent = Inches(-0.25)
        # ----------------------------
        
        # Create numbering properties for bullet
        pPr = para._p.get_or_add_pPr()
        numPr = OxmlElement('w:numPr')
        
        # ilvl = indentation level (0 = first level)
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), '0')
        numPr.append(ilvl)
        
        # numId = numbering definition ID (1 is typically bullet in most templates)
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), '1')
        numPr.append(numId)
        
        pPr.append(numPr)
        
        # Add the text
        run = para.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        
        return para
    
    def _create_normal_paragraph(self, doc, text, font_name="Times New Roman", font_size=10, bold=False):
        """Create a normal paragraph without bullets."""
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold
        return para
    
    def _create_skills_paragraph(self, doc, text, font_name="Times New Roman", font_size=10):
        """Create a skills line with bold category and normal values."""
        para = doc.add_paragraph()
        
        if ":" in text:
            category, skills = text.split(":", 1)
            # Bold category
            run_cat = para.add_run(category + ":")
            run_cat.font.name = font_name
            run_cat.font.size = Pt(font_size)
            run_cat.bold = True
            # Normal skills
            run_skills = para.add_run(skills)
            run_skills.font.name = font_name
            run_skills.font.size = Pt(font_size)
        else:
            run = para.add_run(text)
            run.font.name = font_name
            run.font.size = Pt(font_size)
        
        return para
    
    def _create_env_paragraph(self, doc, text, font_name="Times New Roman", font_size=10):
        """Create an environment line with bold 'Environment:' prefix."""
        para = doc.add_paragraph()
        
        # Ensure it starts with "Environment:" 
        if not text.lower().startswith("environment:"):
            text = "Environment: " + text
        
        if ":" in text:
            prefix, tech_stack = text.split(":", 1)
            # Bold prefix
            run_prefix = para.add_run(prefix + ":")
            run_prefix.font.name = font_name
            run_prefix.font.size = Pt(font_size)
            run_prefix.bold = True
            # Normal tech stack
            run_tech = para.add_run(tech_stack)
            run_tech.font.name = font_name
            run_tech.font.size = Pt(font_size)
        else:
            run = para.add_run(text)
            run.font.name = font_name
            run.font.size = Pt(font_size)
        
        return para

    def fill_template(self, template_path, output_path, final_data):
        """
        Fill the template by finding [[tags]] and replacing them with content.
        Handles tags with surrounding whitespace, tabs, and colons.
        """
        shutil.copy(template_path, output_path)
        doc = Document(output_path)
        
        # Process each tag
        for tag, content_lines in final_data.items():
            
            # --- THE FIX: Strip brackets in case the AI included them in the JSON key ---
            clean_tag = tag.replace("[", "").replace("]", "")
            tag_placeholder = f"[[{clean_tag}]]"
            tag_lower = clean_tag.lower()
            # ----------------------------------------------------------------------------
            
            # Find the paragraph containing this tag
            for i, para in enumerate(doc.paragraphs):
                if tag_placeholder in para.text:
                    # Store the parent element and position
                    parent = para._p.getparent()
                    index = list(parent).index(para._p)
                    
                    # Delete the original tag paragraph entirely
                    parent.remove(para._p)
                    
                    # Insert new paragraphs at the same position
                    paragraphs_to_insert = []
                    
                    for line in content_lines:
                        line = line.strip()
                        if not line:
                            continue
                        
                        # Determine the type of content based on tag name
                        if "description" in tag_lower:
                            new_para = self._create_normal_paragraph(doc, line)
                            paragraphs_to_insert.append(new_para._p)
                            
                        elif "env" in tag_lower:
                            new_para = self._create_env_paragraph(doc, line)
                            paragraphs_to_insert.append(new_para._p)
                            
                        elif "skills" in tag_lower:
                            new_para = self._create_skills_paragraph(doc, line)
                            paragraphs_to_insert.append(new_para._p)
                            
                        else:
                            new_para = self._create_bullet_paragraph(doc, line)
                            paragraphs_to_insert.append(new_para._p)
                    
                    # Now move the paragraphs to the correct position
                    # First, remove them from end of document (where add_paragraph puts them)
                    for p_elem in paragraphs_to_insert:
                        parent_body = p_elem.getparent()
                        if parent_body is not None:
                            parent_body.remove(p_elem)
                    
                    # Insert at the correct position in order
                    for j, p_elem in enumerate(paragraphs_to_insert):
                        parent.insert(index + j, p_elem)
                    
                    break  # Found and processed this tag, move to next
        
        doc.save(output_path)
        print(f"✅ Document saved to: {output_path}")
        return output_path